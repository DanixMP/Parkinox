#!/usr/bin/env python3
"""
Persian (Farsi) Parkinox report builders  -  RTL blue-themed DOCX documents.
Imported and invoked by generate_reports.py.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from generate_reports import (
    BLUE,
    CHARTS,
    ROOT,
    TEHRAN,
    FONT_AR,
    FONT_AR_REG,
    add_chart,
    compute_stats,
    configure_doc_styles,
    fmt_dt,
    fmt_duration,
    get_plate_image,
    save_fig,
    set_cell_shading,
    set_run_font,
    style_axes,
    to_persian_digits,
)

STATUS_FA = {
    "completed": "تکمیل‌شده",
    "waived": "بخشیده‌شده",
    "flagged": "علامت‌خورده",
    "parked": "پارک‌شده",
    "unpaid": "پرداخت‌نشده",
}

PAY_FA = {
    "wallet": "کیف‌پول",
    "cash": "نقدی",
    "waived": "بخشیده",
    "none": "بدون‌پرداخت",
}


def fmt_num(n) -> str:
    try:
        if isinstance(n, float):
            if n != n:  # NaN
                return " - "
            if n == int(n):
                n = int(n)
            else:
                return to_persian_digits(f"{n:,.1f}")
        return to_persian_digits(f"{int(n):,}")
    except Exception:
        return to_persian_digits(str(n))


def fmt_dt_fa(ts) -> str:
    if pd.isna(ts):
        return " - "
    return to_persian_digits(ts.strftime("%Y-%m-%d  %H:%M:%S"))


def fmt_duration_fa(mins) -> str:
    if pd.isna(mins):
        return " - "
    m = int(mins)
    if m < 60:
        return to_persian_digits(f"{m} دقیقه")
    h, rem = divmod(m, 60)
    if h < 24:
        return to_persian_digits(f"{h} ساعت و {rem} دقیقه")
    d, h = divmod(h, 24)
    return to_persian_digits(f"{d} روز و {h} ساعت")


def set_rtl_run(run) -> None:
    r = run._element
    rPr = r.get_or_add_rPr()
    rtl = rPr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        rPr.append(rtl)
    rtl.set(qn("w:val"), "1")
    # Prefer Persian-capable fonts
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), "Noto Sans Arabic")


def set_paragraph_rtl(paragraph, align_right: bool = True) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")
    if align_right:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_doc_rtl(doc: Document) -> None:
    for section in doc.sections:
        sectPr = section._sectPr
        bidi = sectPr.find(qn("w:bidi"))
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            sectPr.append(bidi)
        bidi.set(qn("w:val"), "1")


def fa_font(run, size=11, bold=False, color=None) -> None:
    set_run_font(run, size=size, bold=bold, color=color, name="Noto Sans Arabic")
    set_rtl_run(run)


def add_heading_fa(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    set_paragraph_rtl(p)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(BLUE["navy"].lstrip("#"))
        fa_font(run, size=16 if level == 1 else 13, bold=True, color=BLUE["navy"])


def add_body_fa(doc: Document, text: str, size: int = 11) -> None:
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    run = p.add_run(text)
    fa_font(run, size=size, color=BLUE["text"])
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_bullet_fa(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    set_paragraph_rtl(p)
    run = p.add_run(text)
    fa_font(run, size=11, color=BLUE["text"])


def add_kpi_fa(doc: Document, kpis: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(kpis))
    table.autofit = True
    row = table.rows[0]
    for i, (label, value) in enumerate(kpis):
        cell = row.cells[i]
        set_cell_shading(cell, BLUE["pale"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(value + "\n")
        fa_font(r1, size=16, bold=True, color=BLUE["navy"])
        r2 = p.add_run(label)
        fa_font(r2, size=9, color=BLUE["muted"])
    doc.add_paragraph()


def add_cover_fa(doc: Document, title: str, subtitle: str, meta_lines: list[str], banner: Path) -> None:
    configure_doc_styles(doc)
    set_doc_rtl(doc)
    if banner.exists():
        add_chart(doc, banner, width_in=6.5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    fa_font(r, size=26, bold=True, color=BLUE["navy"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    fa_font(r, size=13, color=BLUE["primary"])

    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        fa_font(r, size=10, color=BLUE["muted"])
    doc.add_paragraph()


def add_caption_fa(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    fa_font(r, size=9, color=BLUE["muted"])


def _setup_fa_matplotlib() -> None:
    from matplotlib import font_manager

    try:
        font_manager.fontManager.addfont(FONT_AR)
        font_manager.fontManager.addfont(FONT_AR_REG)
        font_manager.fontManager.addfont("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
        font_manager.fontManager.addfont("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
    except Exception:
        pass
    # Prefer Persian glyphs; fall back to DejaVu for Latin punctuation/digits in axes
    plt.rcParams["font.family"] = "sans-serif"
    plt.rcParams["font.sans-serif"] = [
        "Noto Sans Arabic",
        "DejaVu Sans",
        "Arial",
        "sans-serif",
    ]
    plt.rcParams["axes.unicode_minus"] = False


def build_charts_fa(df: pd.DataFrame) -> dict[str, Path]:
    """Persian-labeled blue charts (separate files from English)."""
    _setup_fa_matplotlib()
    paths: dict[str, Path] = {}

    daily = df.groupby("entry_date").size()
    fig, ax = plt.subplots(figsize=(11, 4.2))
    ax.fill_between(daily.index, daily.values, color=BLUE["light"], alpha=0.35)
    ax.plot(daily.index, daily.values, color=BLUE["primary"], linewidth=2.2, marker="o", markersize=3.5)
    style_axes(ax, "تعداد نشست‌های پارکینگ در هر روز  -  وقت تهران")
    ax.set_ylabel("نشست")
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%m-%d"))
    fig.autofmt_xdate(rotation=35)
    paths["daily"] = save_fig(fig, "fa_sessions_per_day.png")

    hourly = df.groupby("entry_hour").size().reindex(range(24), fill_value=0)
    fig, ax = plt.subplots(figsize=(10, 4))
    bars = ax.bar(hourly.index, hourly.values, color=BLUE["mid"], edgecolor=BLUE["navy"], width=0.75)
    for i, b in enumerate(bars):
        if hourly.values[i] == hourly.values.max():
            b.set_color(BLUE["navy"])
    style_axes(ax, "ترافیک ورود بر اساس ساعت روز  -  تهران")
    ax.set_xlabel("ساعت")
    ax.set_ylabel("ورود")
    ax.set_xticks(range(0, 24, 2))
    paths["hourly"] = save_fig(fig, "fa_entries_by_hour.png")

    status_counts = df["status"].value_counts()
    labels = [STATUS_FA.get(str(k), str(k)) for k in status_counts.index]
    colors = [BLUE["primary"], BLUE["warn"], BLUE["success"], BLUE["danger"], BLUE["light"]]
    fig, ax = plt.subplots(figsize=(6.5, 6.5))
    wedges, texts, autotexts = ax.pie(
        status_counts.values,
        labels=labels,
        autopct=lambda p: to_persian_digits(f"{p:.1f}٪"),
        colors=colors[: len(status_counts)],
        startangle=90,
        wedgeprops={"linewidth": 2, "edgecolor": "white"},
        textprops={"color": BLUE["text"]},
    )
    for t in autotexts:
        t.set_color("white")
        t.set_fontweight("bold")
    ax.set_title("توزیع وضعیت نشست‌ها", color=BLUE["navy"], fontsize=13, fontweight="bold")
    paths["status"] = save_fig(fig, "fa_status_distribution.png")

    pay = df["payment_method"].value_counts()
    pay_labels = [PAY_FA.get(str(k), str(k)) for k in pay.index]
    fig, ax = plt.subplots(figsize=(8, 4.2))
    ax.barh(pay_labels, pay.values, color=[BLUE["primary"], BLUE["mid"], BLUE["light"], BLUE["sky"]][: len(pay)])
    style_axes(ax, "تفکیک روش پرداخت")
    ax.set_xlabel("نشست")
    for i, v in enumerate(pay.values):
        ax.text(v + max(pay.values) * 0.01, i, fmt_num(v), va="center", color=BLUE["navy"], fontweight="bold")
    paths["payment"] = save_fig(fig, "fa_payment_methods.png")

    closed = df[df["duration_minutes"].notna()].copy()
    bins = [0, 5, 15, 30, 60, 120, 240, 10**9]
    labels_d = ["۰ تا ۵د", "۵ تا ۱۵د", "۱۵ تا ۳۰د", "۳۰ تا ۶۰د", "۱ تا ۲س", "۲ تا ۴س", "بیش از ۴س"]
    closed["bucket"] = pd.cut(closed["duration_minutes"], bins=bins, labels=labels_d, right=False)
    bucket_counts = closed["bucket"].value_counts().reindex(labels_d).fillna(0)
    fig, ax = plt.subplots(figsize=(9, 4.2))
    ax.bar(labels_d, bucket_counts.values, color=BLUE["accent"], edgecolor=BLUE["navy"])
    style_axes(ax, "توزیع مدت توقف")
    ax.set_ylabel("نشست")
    paths["duration"] = save_fig(fig, "fa_duration_distribution.png")

    daily_s = df.groupby("entry_date").size()
    top_days = daily_s.sort_index().tail(14).index
    sub = df[df["entry_date"].isin(top_days)]
    pivot = pd.crosstab(sub["entry_date"], sub["status"]).reindex(top_days).fillna(0)
    fig, ax = plt.subplots(figsize=(11, 4.5))
    bottom = None
    palette = {
        "waived": BLUE["light"],
        "flagged": BLUE["warn"],
        "completed": BLUE["success"],
        "parked": BLUE["primary"],
        "unpaid": BLUE["danger"],
    }
    for col in pivot.columns:
        vals = pivot[col].values
        ax.bar(
            range(len(pivot)),
            vals,
            bottom=bottom,
            label=STATUS_FA.get(str(col), str(col)),
            color=palette.get(str(col), BLUE["mid"]),
            edgecolor="white",
            width=0.7,
        )
        bottom = vals if bottom is None else bottom + vals
    ax.set_xticks(range(len(pivot)))
    ax.set_xticklabels([d.strftime("%m-%d") for d in pivot.index], rotation=35)
    style_axes(ax, "ترکیب وضعیت روزانه  -  روزهای فعال اخیر")
    ax.legend(frameon=False, loc="upper left")
    ax.set_ylabel("نشست")
    paths["status_daily"] = save_fig(fig, "fa_status_by_day.png")

    fee_pos = df[df["fee"] > 0]["fee"]
    fig, ax = plt.subplots(figsize=(9, 4.2))
    if len(fee_pos):
        ax.hist(fee_pos, bins=30, color=BLUE["mid"], edgecolor=BLUE["navy"], alpha=0.9)
    style_axes(ax, "توزیع مبلغ محاسبه‌شده  -  تومان، مبلغ بیشتر از صفر")
    ax.set_xlabel("مبلغ (تومان)")
    ax.set_ylabel("نشست")
    paths["fees"] = save_fig(fig, "fa_fee_distribution.png")

    cam = pd.Series(
        {
            "فقط ورود": int(((df["entry_camera"] != "") & (df["exit_camera"] == "")).sum()),
            "ورود + خروج": int(((df["entry_camera"] != "") & (df["exit_camera"] != "")).sum()),
            "بدون برچسب دوربین": int(((df["entry_camera"] == "") & (df["exit_camera"] == "")).sum()),
        }
    )
    fig, ax = plt.subplots(figsize=(7, 4.2))
    ax.bar(cam.index, cam.values, color=[BLUE["primary"], BLUE["success"], BLUE["muted"]])
    style_axes(ax, "پوشش برچسب دوربین روی نشست‌ها")
    ax.set_ylabel("نشست")
    paths["cameras"] = save_fig(fig, "fa_camera_coverage.png")

    fig, ax = plt.subplots(figsize=(12, 2.2))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.imshow([[0, 1], [0, 1]], extent=[0, 1, 0, 1], aspect="auto", cmap="Blues", alpha=0.95)
    ax.text(0.5, 0.58, "پارکینوکس", color="white", fontsize=28, fontweight="bold", ha="center", va="center")
    ax.text(
        0.5,
        0.28,
        "پارکینگ هوشمند دانشگاهی  -  گزارش موفقیت و تحلیل داده",
        color="#E3F2FD",
        fontsize=12,
        ha="center",
    )
    paths["banner"] = save_fig(fig, "fa_banner.png")

    return paths


def build_success_report_fa(df: pd.DataFrame, charts: dict[str, Path], stats: dict) -> Path:
    doc = Document()
    add_cover_fa(
        doc,
        "گزارش موفقیت پارکینوکس",
        "پلتفرم یکپارچه پارکینگ هوشمند  -  موفقیت عملیاتی و دفترچه نشست‌ها",
        [
            "نسخه تم آبی · تشخیص پلاک ایرانی · معماری محلی‌محور",
            f"بازه داده: {fmt_dt_fa(stats['date_min'])}  ←  {fmt_dt_fa(stats['date_max'])} (تهران)",
            f"تاریخ تولید: {to_persian_digits(datetime.now(TEHRAN).strftime('%Y-%m-%d %H:%M'))} Asia/Tehran",
            "منبع: parkinox_sessions.csv · پروژه: Parkinox v3",
        ],
        charts["banner"],
    )

    add_heading_fa(doc, "۱. خلاصه مدیریتی", 1)
    add_body_fa(
        doc,
        "پارکینوکس نسخه ۳ با موفقیت یک پلتفرم پارکینگ دانشگاهی حافظ حریم خصوصی را ارائه کرد که "
        "تشخیص محلی پلاک ایرانی (YOLOv5 روی رایانه اپراتور)، منطق کسب‌وکار جنگو برای نشست‌ها/تعرفه/"
        "کیف‌پول، کلاینت‌های فلاتر اپراتور و دانشجو، و همگام‌سازی پایدار Outbox به ابر گزارش‌گیری را "
        f"ترکیب می‌کند. داده عملیاتی پیوست، {fmt_num(stats['total'])} نشست پارکینگ برای "
        f"{fmt_num(stats['unique_plates'])} پلاک یکتا را نشان می‌دهد؛ تأیید گیت، محاسبه تعرفه، مسیرهای "
        "پرداخت، و گردش‌کار بخشش/علامت‌گذاری اپراتور مطابق طراحی عمل کرده‌اند.",
    )
    add_kpi_fa(
        doc,
        [
            ("کل نشست‌ها", fmt_num(stats["total"])),
            ("پلاک یکتا", fmt_num(stats["unique_plates"])),
            ("تکمیل‌شده", fmt_num(stats["completed"])),
            ("بخشیده‌شده", fmt_num(stats["waived"])),
            ("علامت‌خورده / باز", fmt_num(stats["flagged"])),
        ],
    )

    add_heading_fa(doc, "۲. نحوه عملکرد سامانه", 1)
    add_heading_fa(doc, "۲.۱ معماری", 2)
    add_body_fa(
        doc,
        "تشخیص روی ماشین اپراتور می‌ماند: FastAPI به‌همراه OpenCV و YOLOv5 جریان دوربین RTSP/USB را "
        "می‌گیرد، رویداد plate_detected را منتشر می‌کند و پیش‌نمایش MJPEG می‌دهد. سرور جنگوی راه‌دور "
        "منبع حقیقت برای احراز هویت، نشست پارکینگ، رویداد گیت، کیف‌پول و گزارش است. دستگاه‌های دانشجو "
        "از اپلیکیشن موبایل فلاتر (OTP، ثبت پلاک، کیف‌پول) استفاده می‌کنند. یک نسخه ابری جداگانه از "
        "جنگو، فراداده و تصویر را از طریق SyncOutbox برای تحلیل دریافت می‌کند  -  هرگز ویدئوی خام دوربین را.",
    )
    add_bullet_fa(doc, "parkinox_op  -  رابط دسکتاپ اپراتور فلاتر (تأیید، گیت، بازبینی خطا)")
    add_bullet_fa(doc, "Core/FastAPI  -  مدل‌های plateYolo و CharsYolo، اعتبارسنجی، ثبت موفقیت/خطا")
    add_bullet_fa(doc, "backend/Django  -  GateEvent، ParkingSession، RateConfig، کیف‌پول، گزارش‌ها")
    add_bullet_fa(doc, "parkinox_stdn  -  دسترسی موبایل دانشجو")
    add_bullet_fa(doc, "cloud/  -  داشبورد گزارش فارسی RTL (Chart.js، اکسل/PDF/CSV)")

    add_heading_fa(doc, "۲.۲ چرخه حیات نشست", 2)
    add_body_fa(
        doc,
        "فریم دوربین ← محدوده پلاک YOLO ← OCR کاراکتر ← اعتبارسنجی قالب ایرانی (همان نرمالایزر جنگو) "
        "← شمارش معکوس تأیید خودکار یا تأیید دستی اپراتور ← confirm_gate_event():",
    )
    add_bullet_fa(doc, "ورود: ایجاد ParkingSession با وضعیت parked؛ اتصال GateEvent تغییرناپذیر")
    add_bullet_fa(doc, "خروج: بستن نشست، محاسبه مدت/مبلغ (مهلت رایگان، نرخ ساعتی، سقف روزانه، تخفیف دانشجو)")
    add_bullet_fa(doc, "پرداخت: کسر خودکار کیف‌پول در صورت ثبت‌نام و موجودی کافی؛ در غیر این صورت unpaid ← نقدی یا بخشش")
    add_bullet_fa(doc, "پس از commit: صف SyncOutbox + نهایی‌سازی success-capture (صحنه/برش)")

    add_heading_fa(doc, "۲.۳ مدیریت پلاک ایرانی", 2)
    add_body_fa(
        doc,
        "پلاک ملی به‌صورت کانونیکال ۱۲ب۳۴۵-۶۷ و پلاک منطقه آزاد با ۵+۲ رقم است. نرمال‌سازی ارقام "
        "عربی/لاتین را به فارسی تبدیل می‌کند، نویز را حذف و مجموعه حروف مجاز را بررسی می‌کند. نمایش "
        "اپراتور پلاک را فاصله‌گذاری می‌کند (۱۲ ب ۳۴۵ ۶۷). محافظ بین‌دوربینی جهت مخالف همان پلاک را "
        "به‌مدت دو دقیقه مسدود می‌کند تا خطای برگشت سریع رخ ندهد.",
    )

    sample = df.head(6)
    add_body_fa(doc, "نمونه‌هایی از پلاک‌ها با سبک پلاک ایرانی (از خروجی عملیاتی):")
    table = doc.add_table(rows=1, cols=3)
    for i, (_, row) in enumerate(sample.iterrows()):
        if i >= 3:
            break
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, BLUE["pale"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(get_plate_image(row["plate"])), width=Inches(1.85))
    doc.add_paragraph()

    add_heading_fa(doc, "۳. مدیریت خطا و تاب‌آوری", 1)
    add_heading_fa(doc, "۳.۱ ثبت خطای تشخیص", 2)
    add_body_fa(doc, "OCR ناموفق هرگز گیت را به‌صورت خودکار باز نمی‌کند. خطاها طبقه‌بندی و برای بازبینی اپراتور ذخیره می‌شوند:")
    add_bullet_fa(doc, "نوع A  -  ocr_invalid_format: محدوده پیدا شد، متن قواعد ایرانی را پاس نکرد")
    add_bullet_fa(doc, "نوع B  -  ocr_empty: محدوده پیدا شد، OCR خالی یا با اطمینان پایین")
    add_bullet_fa(doc, "نوع C  -  operator_confirmed_miss: پلاک دیده می‌شود اما محدوده نیست (ثبت اپراتور)")
    add_bullet_fa(doc, "نوع D  -  معیار موفقیت تلاش مجدد گذرا (خطا سپس تشخیص معتبر طی حدود ۵ ثانیه)")
    add_body_fa(
        doc,
        "FailCaptureStore در Core شامل کول‌داون، محدودیت نرخ هر دوربین، حذف تکراری با IoU و انتخاب "
        "بهترین فریم در پنجره زمانی است و سپس به DetectionFail جنگو وارد می‌شود. اپراتور می‌تواند پلاک "
        "را اصلاح کند (با فراخوانی confirm_gate_event و زمان captured_at اصلی) یا not_a_plate علامت بزند.",
    )

    add_heading_fa(doc, "۳.۲ یکپارچگی گیت و نشست", 2)
    add_bullet_fa(doc, "GateEvent یک لاگ حسابرسی تغییرناپذیر است (فقط درج)")
    add_bullet_fa(doc, "پرچم‌ها: was_edited، was_auto_approved، was_rejected")
    add_bullet_fa(doc, "نشست‌های همزمان parked برای یک پلاک قبل از ورود مجدد علامت/بسته می‌شوند")
    add_bullet_fa(doc, "اصلاح خطای تاریخی اگر نشست زنده جدیدتری را خراب کند رد می‌شود")
    add_bullet_fa(doc, "کسر کیف‌پول با select_for_update؛ موجودی ناکافی ← unpaid (بدون کسر جزئی)")

    add_heading_fa(doc, "۳.۳ تاب‌آوری خط لوله و همگام‌سازی", 2)
    add_bullet_fa(doc, "کارگرهای دوربین: backoff اتصال مجدد، فیلتر تاری، فاصله تشخیص، کول‌داون ارسال")
    add_bullet_fa(doc, "ثبت موفقیت فریم‌ها را بافر می‌کند؛ finalize مسدودکننده نیست")
    add_bullet_fa(doc, "SyncOutbox فقط پس از commit پایگاه‌داده صف می‌شود  -  هرگز در مسیر بحرانی گیت")
    add_bullet_fa(doc, "وضعیت Outbox: pending ← sending ← acked / dead_letter (حداکثر ۱۲ تلاش)")
    add_bullet_fa(doc, "گیت و گزارش محلی هنگام قطع ابر همچنان کار می‌کنند")

    add_heading_fa(doc, "۴. نتایج عملیاتی از داده نشست", 1)
    peak_hour_fa = to_persian_digits(f"{stats['peak_hour']:02d}:00")
    add_body_fa(
        doc,
        f"در بازه خروجی، روز اوج {to_persian_digits(str(stats['peak_day']))} با "
        f"{fmt_num(stats['peak_day_count'])} نشست ثبت شد. ساعت اوج ورود (تهران) "
        f"{peak_hour_fa} بود. "
        f"{fmt_num(stats['with_exit'])} نشست دارای زمان خروج هستند؛ "
        f"{fmt_num(stats['open_or_flagged_no_exit'])} بدون خروج مانده‌اند (معمولاً علامت‌خورده). "
        f"روش‌های پرداخت مشاهده‌شده: کیف‌پول={fmt_num(stats['wallet'])}، نقدی={fmt_num(stats['cash'])}. "
        f"گردش‌کار بخشش اپراتور {fmt_num(stats['waived'])} بستن را پوشش می‌دهد  -  هم‌راستا با سیاست "
        "پایلوت/پردیس و نه نقص سامانه.",
    )
    add_kpi_fa(
        doc,
        [
            ("میانگین مدت", fmt_duration_fa(stats["avg_duration"])),
            ("میانه مدت", fmt_duration_fa(stats["median_duration"])),
            ("جمع مبلغ (تومان)", fmt_num(stats["fee_sum"])),
            ("کاربر نام‌دار", fmt_num(stats["registered_named"])),
            ("ساعت اوج", peak_hour_fa),
        ],
    )

    add_heading_fa(doc, "۵. نمودارهای تحلیلی", 1)
    add_body_fa(doc, "نمودارهای زیر از همان CSV گزارش تحلیلی همراه ساخته شده‌اند.")
    for key, caption in [
        ("daily", "شکل ۱  -  نشست در هر روز"),
        ("hourly", "شکل ۲  -  ترافیک ورود بر اساس ساعت"),
        ("status", "شکل ۳  -  توزیع وضعیت"),
        ("payment", "شکل ۴  -  روش‌های پرداخت"),
        ("duration", "شکل ۵  -  بازه‌های مدت توقف"),
        ("status_daily", "شکل ۶  -  ترکیب وضعیت روزانه اخیر"),
        ("fees", "شکل ۷  -  توزیع مبلغ"),
        ("cameras", "شکل ۸  -  پوشش برچسب دوربین"),
    ]:
        add_chart(doc, charts[key])
        add_caption_fa(doc, caption)

    add_heading_fa(doc, "۶. دفترچه کامل نشست‌ها  -  همه پلاک‌ها", 1)
    add_body_fa(
        doc,
        "همه نشست‌های CSV، مرتب‌شده بر اساس زمان ورود (جدیدترین ابتدا). هر ردیف نشان پلاک ایرانی، "
        "وضعیت، زمان ورود/خروج تهران، مدت، مبلغ و روش پرداخت را نمایش می‌دهد.",
    )

    header_labels = ["#", "پلاک", "وضعیت", "ورود (تهران)", "خروج (تهران)", "مدت", "مبلغ", "پرداخت"]
    batch_size = 40
    total = len(df)
    for start in range(0, total, batch_size):
        chunk = df.iloc[start : start + batch_size]
        end = min(start + batch_size, total)
        add_heading_fa(doc, f"نشست‌های {fmt_num(start + 1)} تا {fmt_num(end)}", 2)

        table = doc.add_table(rows=1, cols=len(header_labels))
        table.style = "Table Grid"
        hdr = table.rows[0]
        for i, label in enumerate(header_labels):
            cell = hdr.cells[i]
            set_cell_shading(cell, BLUE["navy"])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(label)
            fa_font(r, size=8, bold=True, color=BLUE["white"])

        for idx, (_, row) in enumerate(chunk.iterrows(), start=start + 1):
            cells = table.add_row().cells
            values = [
                fmt_num(idx),
                "",
                STATUS_FA.get(str(row["status"]), str(row["status"])),
                fmt_dt_fa(row["entry_tehran"]),
                fmt_dt_fa(row["exit_tehran"]),
                fmt_duration_fa(row["duration_minutes"]),
                fmt_num(int(row["fee"])) if row["fee"] else to_persian_digits("0"),
                PAY_FA.get(str(row["payment_method"]), str(row["payment_method"])),
            ]
            for i, val in enumerate(values):
                p = cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i == 1:
                    run = p.add_run()
                    run.add_picture(str(get_plate_image(row["plate"])), width=Inches(1.35))
                else:
                    r = p.add_run(val)
                    fa_font(r, size=7, color=BLUE["text"])
                if idx % 2 == 0:
                    set_cell_shading(cells[i], BLUE["pale"])
        doc.add_paragraph()

    add_heading_fa(doc, "۷. نتیجه‌گیری", 1)
    add_body_fa(
        doc,
        "پارکینوکس معیارهای موفقیت پایلوت را برآورده کرد: تشخیص محلی پلاک ایرانی با OCR معتبر، "
        "تأیید گیت تحت کنترل اپراتور، حسابداری پایدار نشست، مسیرهای پرداخت کیف‌پول/نقدی/بخشش، "
        "ثبت ساختاریافته خطا بدون باز شدن ناامن خودکار گیت، و گزارش ابری که منبع حقیقت محلی را "
        "تضعیف نمی‌کند. دفترچه نشست‌ها و نمودارهای این سند، سابقه قابل حسابرسی از رفتار سامانه "
        "در ترافیک واقعی پردیس را فراهم می‌کنند.",
    )

    out = ROOT / "Parkinox_Success_Report_FA.docx"
    doc.save(out)
    return out


def build_data_report_fa(df: pd.DataFrame, charts: dict[str, Path], stats: dict) -> Path:
    doc = Document()
    add_cover_fa(
        doc,
        "گزارش تحلیل داده پارکینوکس",
        "نمودارهای مجموعه داده نشست و بینش‌های اندازه‌گیری‌شده زمان",
        [
            "تم آبی · سند صرفاً مبتنی بر داده (CSV)",
            f"ردیف‌ها: {fmt_num(stats['total'])}  ·  پلاک یکتا: {fmt_num(stats['unique_plates'])}",
            f"بازه: {fmt_dt_fa(stats['date_min'])}  ←  {fmt_dt_fa(stats['date_max'])} (تهران)",
            f"تاریخ تولید: {to_persian_digits(datetime.now(TEHRAN).strftime('%Y-%m-%d %H:%M'))} Asia/Tehran",
            "فایل منبع: parkinox_sessions.csv",
        ],
        charts["banner"],
    )

    add_heading_fa(doc, "۱. نمای کلی مجموعه داده", 1)
    add_body_fa(
        doc,
        "این گزارش فقط اندازه‌گیری‌های مشتق‌شده از CSV نشست‌های پارکینگ را شامل می‌شود. "
        "برچسب‌های زمانی از UTC به Asia/Tehran (UTC+۰۳:۳۰) برای تحلیل ترافیک تبدیل شده‌اند. "
        "فیلدهای مدت و مبلغ مطابق خروجی هستند؛ نشست‌های بدون exit_time به‌عنوان باز/علامت‌خورده در نظر گرفته می‌شوند.",
    )
    add_kpi_fa(
        doc,
        [
            ("نشست‌ها", fmt_num(stats["total"])),
            ("پلاک یکتا", fmt_num(stats["unique_plates"])),
            ("دارای خروج", fmt_num(stats["with_exit"])),
            ("بدون خروج", fmt_num(stats["open_or_flagged_no_exit"])),
            ("جمع مبلغ", fmt_num(stats["fee_sum"])),
        ],
    )

    add_heading_fa(doc, "۲. ساختار فیلدها", 1)
    schema = [
        ("plate", "رشته پلاک ایرانی (ارقام فارسی + حرف)"),
        ("status", "تکمیل‌شده | بخشیده‌شده | علامت‌خورده"),
        ("entry_time / exit_time", "برچسب زمانی ISO-8601 با UTC"),
        ("duration_minutes", "مدت توقف هنگام خروج"),
        ("fee", "مبلغ محاسبه‌شده به تومان"),
        ("payment_method", "کیف‌پول | نقدی | بخشیده | بدون‌پرداخت"),
        ("entry_camera / exit_camera", "برچسب دوربین در صورت وجود"),
        ("user_name", "نام نمایشی کاربر ثبت‌شده"),
        ("session_uuid", "شناسه پایدار نشست"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, h in enumerate(["فیلد", "معنا"]):
        set_cell_shading(table.rows[0].cells[i], BLUE["navy"])
        r = table.rows[0].cells[i].paragraphs[0].add_run(h)
        fa_font(r, size=10, bold=True, color=BLUE["white"])
        table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for field, meaning in schema:
        row = table.add_row().cells
        r0 = row[0].paragraphs[0].add_run(field)
        fa_font(r0, size=9, bold=True, color=BLUE["navy"])
        r1 = row[1].paragraphs[0].add_run(meaning)
        fa_font(r1, size=9, color=BLUE["text"])
        set_paragraph_rtl(row[1].paragraphs[0])
    doc.add_paragraph()

    add_heading_fa(doc, "۳. شمارش وضعیت و پرداخت", 1)
    status = df["status"].value_counts()
    pay = df["payment_method"].value_counts()
    add_body_fa(doc, "شمارش وضعیت:")
    for k, v in status.items():
        add_bullet_fa(
            doc,
            f"{STATUS_FA.get(str(k), str(k))}: {fmt_num(v)} ({to_persian_digits(f'{100 * v / len(df):.1f}')}٪)",
        )
    add_body_fa(doc, "شمارش روش پرداخت:")
    for k, v in pay.items():
        add_bullet_fa(
            doc,
            f"{PAY_FA.get(str(k), str(k))}: {fmt_num(v)} ({to_persian_digits(f'{100 * v / len(df):.1f}')}٪)",
        )

    add_heading_fa(doc, "۴. نمودارهای سری زمانی", 1)
    peak_hour_fa = to_persian_digits(f"{stats['peak_hour']:02d}:00")
    add_body_fa(
        doc,
        f"روز اوج: {to_persian_digits(str(stats['peak_day']))} با {fmt_num(stats['peak_day_count'])} ورود. "
        f"ساعت اوج (تهران): {peak_hour_fa}. "
        f"میانگین مدت نشست‌های بسته‌شده: {fmt_duration_fa(stats['avg_duration'])}؛ "
        f"میانه: {fmt_duration_fa(stats['median_duration'])}.",
    )
    for key, caption in [
        ("daily", "نمودار الف  -  نشست در هر روز"),
        ("hourly", "نمودار ب  -  ورود بر اساس ساعت"),
        ("status_daily", "نمودار ج  -  ترکیب وضعیت روزهای فعال اخیر"),
    ]:
        add_chart(doc, charts[key])
        add_caption_fa(doc, caption)

    add_heading_fa(doc, "۵. نمودارهای ترکیبی", 1)
    for key, caption in [
        ("status", "نمودار د  -  توزیع وضعیت"),
        ("payment", "نمودار هـ  -  روش‌های پرداخت"),
        ("duration", "نمودار و  -  توزیع مدت"),
        ("fees", "نمودار ز  -  توزیع مبلغ (مبلغ > ۰)"),
        ("cameras", "نمودار ح  -  پوشش برچسب دوربین"),
    ]:
        add_chart(doc, charts[key])
        add_caption_fa(doc, caption)

    add_heading_fa(doc, "۶. جدول ترافیک روزانه", 1)
    daily = df.groupby("entry_date").agg(
        sessions=("plate", "size"),
        unique_plates=("plate", "nunique"),
        completed=("status", lambda s: (s == "completed").sum()),
        waived=("status", lambda s: (s == "waived").sum()),
        flagged=("status", lambda s: (s == "flagged").sum()),
        fee_sum=("fee", "sum"),
    )
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["تاریخ", "نشست", "یکتا", "تکمیل", "بخشش", "علامت", "جمع مبلغ"]
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], BLUE["primary"])
        r = table.rows[0].cells[i].paragraphs[0].add_run(h)
        fa_font(r, size=8, bold=True, color=BLUE["white"])
        table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for date, row in daily.iterrows():
        cells = table.add_row().cells
        vals = [
            to_persian_digits(str(date)),
            fmt_num(int(row["sessions"])),
            fmt_num(int(row["unique_plates"])),
            fmt_num(int(row["completed"])),
            fmt_num(int(row["waived"])),
            fmt_num(int(row["flagged"])),
            fmt_num(int(row["fee_sum"])),
        ]
        for i, v in enumerate(vals):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(v)
            fa_font(r, size=8, color=BLUE["text"])
    doc.add_paragraph()

    add_heading_fa(doc, "۷. جدول ترافیک ساعتی", 1)
    hourly = df.groupby("entry_hour").size().reindex(range(24), fill_value=0)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, h in enumerate(["ساعت (تهران)", "ورود"]):
        set_cell_shading(table.rows[0].cells[i], BLUE["primary"])
        r = table.rows[0].cells[i].paragraphs[0].add_run(h)
        fa_font(r, size=9, bold=True, color=BLUE["white"])
        table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for hour, count in hourly.items():
        cells = table.add_row().cells
        for i, v in enumerate([to_persian_digits(f"{int(hour):02d}:00"), fmt_num(int(count))]):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(v)
            fa_font(r, size=9, color=BLUE["text"])
            if count == hourly.max() and count > 0:
                set_cell_shading(cells[i], BLUE["sky"])
    doc.add_paragraph()

    add_heading_fa(doc, "۸. آمار مدت و مبلغ", 1)
    closed = df[df["duration_minutes"].notna()]
    add_bullet_fa(doc, f"نشست‌های بسته‌شده: {fmt_num(len(closed))}")
    add_bullet_fa(doc, f"میانگین مدت: {fmt_duration_fa(closed['duration_minutes'].mean())}")
    add_bullet_fa(doc, f"میانه مدت: {fmt_duration_fa(closed['duration_minutes'].median())}")
    add_bullet_fa(doc, f"صدک نودم مدت: {fmt_duration_fa(closed['duration_minutes'].quantile(0.9))}")
    add_bullet_fa(doc, f"حداکثر مدت: {fmt_duration_fa(closed['duration_minutes'].max())}")
    add_bullet_fa(doc, f"جمع مبالغ محاسبه‌شده: {fmt_num(stats['fee_sum'])} تومان")
    add_bullet_fa(doc, f"میانگین مبلغ در موارد مبلغ > ۰: {fmt_num(round(stats['fee_mean_pos']))} تومان")
    add_bullet_fa(doc, f"نشست با مبلغ صفر: {fmt_num(int((df['fee'] == 0).sum()))}")

    add_heading_fa(doc, "۹. نکات اندازه‌گیری", 1)
    add_body_fa(
        doc,
        "لحظات ورود و خروج از برچسب‌های زمانی ISO با اختلاف ساعت پارس شده و برای تجمیع ساعتی/روزانه "
        "به وقت محلی تهران تبدیل می‌شوند. مدت از فیلد duration_minutes فایل CSV گرفته شده (نه محاسبه "
        "دوباره) تا با خروجی هم‌خوان بماند. وضعیت و روش پرداخت به‌صورت شمارش طبقه‌ای هستند؛ مقادیر خالی "
        "روش پرداخت با برچسب بدون‌پرداخت مشخص شده‌اند. این سند عمداً روایت محصول را ندارد  -  "
        "برای عملکرد سامانه و مدیریت خطا به Parkinox_Success_Report_FA.docx مراجعه کنید.",
    )

    out = ROOT / "Parkinox_Data_Analytics_Report_FA.docx"
    doc.save(out)
    return out
