#!/usr/bin/env python3
"""
Parkinox Success Report generator.

Produces (English + Persian):
  1) Parkinox_Success_Report.docx / _FA.docx
  2) Parkinox_Data_Analytics_Report.docx / _FA.docx
"""

from __future__ import annotations

import hashlib
import re
from datetime import datetime, timezone, timedelta
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths & theme
# ---------------------------------------------------------------------------

ROOT = Path(__file__).resolve().parent
CSV_PATH = ROOT / "parkinox_sessions.csv"
CHARTS = ROOT / "charts"
PLATES = ROOT / "plates"
ASSETS = ROOT / "assets"

BLUE = {
    "navy": "#0B3A6E",
    "primary": "#1565C0",
    "mid": "#1E88E5",
    "light": "#42A5F5",
    "pale": "#E3F2FD",
    "sky": "#BBDEFB",
    "accent": "#0277BD",
    "text": "#0D2137",
    "muted": "#546E7A",
    "white": "#FFFFFF",
    "success": "#2E7D32",
    "warn": "#F9A825",
    "danger": "#C62828",
}

TEHRAN = timezone(timedelta(hours=3, minutes=30))

FONT_AR = "/usr/share/fonts/truetype/noto/NotoSansArabic-Bold.ttf"
FONT_AR_REG = "/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf"
FONT_LATIN = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
FONT_LATIN_REG = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"

NATIONAL_RE = re.compile(r"^([۰-۹0-9]{2})([ء-یآا]{1,4})([۰-۹0-9]{3})-([۰-۹0-9]{2})$")
FREE_ZONE_RE = re.compile(r"^([۰-۹0-9]{5})-([۰-۹0-9]{2})$")

DIGIT_MAP = str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹")


def to_persian_digits(s: str) -> str:
    return str(s).translate(DIGIT_MAP)


def ensure_dirs() -> None:
    for d in (CHARTS, PLATES, ASSETS):
        d.mkdir(parents=True, exist_ok=True)


def load_data() -> pd.DataFrame:
    df = pd.read_csv(CSV_PATH)
    df["entry_time"] = pd.to_datetime(df["entry_time"], utc=True)
    df["exit_time"] = pd.to_datetime(df["exit_time"], utc=True, errors="coerce")
    df["entry_tehran"] = df["entry_time"].dt.tz_convert(TEHRAN)
    df["exit_tehran"] = df["exit_time"].dt.tz_convert(TEHRAN)
    df["entry_date"] = df["entry_tehran"].dt.date
    df["entry_hour"] = df["entry_tehran"].dt.hour
    df["duration_minutes"] = pd.to_numeric(df["duration_minutes"], errors="coerce")
    df["fee"] = pd.to_numeric(df["fee"], errors="coerce").fillna(0)
    df["payment_method"] = df["payment_method"].fillna("none")
    df["exit_camera"] = df["exit_camera"].fillna("")
    df["user_name"] = df["user_name"].fillna("")
    # Sort: newest first, then plate
    df = df.sort_values(["entry_time", "plate"], ascending=[False, True]).reset_index(drop=True)
    return df


def fmt_dt(ts) -> str:
    if pd.isna(ts):
        return "—"
    return ts.strftime("%Y-%m-%d  %H:%M:%S")


def fmt_duration(mins) -> str:
    if pd.isna(mins):
        return "—"
    m = int(mins)
    if m < 60:
        return f"{m} min"
    h, rem = divmod(m, 60)
    if h < 24:
        return f"{h}h {rem}m"
    d, h = divmod(h, 24)
    return f"{d}d {h}h"


# ---------------------------------------------------------------------------
# Iranian plate renderer
# ---------------------------------------------------------------------------

def _font(path: str, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(path, size)


def parse_plate(plate: str) -> dict | None:
    plate = str(plate).strip()
    m = NATIONAL_RE.match(plate)
    if m:
        return {
            "kind": "national",
            "left": to_persian_digits(m.group(1)),
            "letter": m.group(2),
            "mid": to_persian_digits(m.group(3)),
            "region": to_persian_digits(m.group(4)),
            "raw": plate,
        }
    m = FREE_ZONE_RE.match(plate)
    if m:
        return {
            "kind": "freezone",
            "serial": to_persian_digits(m.group(1)),
            "region": to_persian_digits(m.group(2)),
            "raw": plate,
        }
    return {"kind": "raw", "raw": to_persian_digits(plate)}


def render_iranian_plate(plate: str, width: int = 320, height: int = 72) -> Image.Image:
    """Draw a realistic Iranian national / free-zone plate badge."""
    parsed = parse_plate(plate)
    img = Image.new("RGB", (width, height), BLUE["white"])
    draw = ImageDraw.Draw(img)

    # Outer black border
    draw.rounded_rectangle([1, 1, width - 2, height - 2], radius=6, outline="#111111", width=2)

    # Left blue IR strip
    strip_w = int(width * 0.14)
    draw.rounded_rectangle([2, 2, strip_w, height - 3], radius=4, fill=BLUE["primary"])
    # Fake flag bands
    band_h = max(3, height // 12)
    fy = 8
    for color in ("#239F40", "#FFFFFF", "#DA0000"):
        draw.rectangle([strip_w // 4, fy, strip_w - strip_w // 4, fy + band_h], fill=color)
        fy += band_h + 1
    try:
        f_ir = _font(FONT_LATIN, max(9, height // 7))
        f_ir_sm = _font(FONT_LATIN, max(7, height // 9))
    except OSError:
        f_ir = ImageFont.load_default()
        f_ir_sm = f_ir
    draw.text((strip_w // 2, height - 16), "I.R.", font=f_ir_sm, fill="white", anchor="mm")
    draw.text((strip_w // 2, height - 6), "IRAN", font=f_ir, fill="white", anchor="mm")

    # Right region box
    region_w = int(width * 0.18)
    rx0 = width - region_w - 3
    draw.rectangle([rx0, 3, width - 3, height - 3], fill=BLUE["pale"], outline=BLUE["mid"])

    body_left = strip_w + 6
    body_right = rx0 - 6
    body_cx = (body_left + body_right) // 2
    cy = height // 2

    try:
        f_big = _font(FONT_AR, max(22, int(height * 0.48)))
        f_letter = _font(FONT_AR, max(20, int(height * 0.42)))
        f_reg = _font(FONT_AR, max(20, int(height * 0.42)))
        f_raw = _font(FONT_AR, max(16, int(height * 0.35)))
    except OSError:
        f_big = f_letter = f_reg = f_raw = ImageFont.load_default()

    # Draw segments at fixed LTR x positions so Arabic shaping cannot reverse order.
    region_cx = (rx0 + width - 3) // 2

    def draw_ltr_group(text: str, cx: int, font, fill: str, y: int = cy) -> None:
        """Draw characters left-to-right around a center x."""
        if not text:
            return
        # Measure individual glyphs
        widths = []
        for ch in text:
            bbox = draw.textbbox((0, 0), ch, font=font)
            widths.append(bbox[2] - bbox[0])
        gap = max(1, int(height * 0.02))
        total = sum(widths) + gap * (len(text) - 1)
        x = cx - total // 2
        for ch, w in zip(text, widths):
            draw.text((x + w // 2, y), ch, font=font, fill=fill, anchor="mm")
            x += w + gap

    if parsed["kind"] == "national":
        span = body_right - body_left
        x1 = body_left + int(span * 0.18)
        x2 = body_left + int(span * 0.48)
        x3 = body_left + int(span * 0.78)
        draw_ltr_group(parsed["left"], x1, f_big, "#111111")
        draw.text((x2, cy), parsed["letter"], font=f_letter, fill="#111111", anchor="mm")
        draw_ltr_group(parsed["mid"], x3, f_big, "#111111")
        draw_ltr_group(parsed["region"], region_cx, f_reg, BLUE["navy"])
    elif parsed["kind"] == "freezone":
        draw_ltr_group(parsed["serial"], body_cx, f_big, "#111111", y=cy - 6)
        draw.text((body_cx, cy + 18), "FREE ZONE", font=f_ir_sm, fill=BLUE["accent"], anchor="mm")
        draw_ltr_group(parsed["region"], region_cx, f_reg, BLUE["navy"])
    else:
        draw.text((body_cx, cy), parsed["raw"], font=f_raw, fill="#111111", anchor="mm")
        draw.text((region_cx, cy), "—", font=f_reg, fill=BLUE["muted"], anchor="mm")

    return img


def plate_path(plate: str) -> Path:
    key = hashlib.md5(str(plate).encode("utf-8")).hexdigest()[:16]
    return PLATES / f"plate_{key}.png"


def get_plate_image(plate: str) -> Path:
    path = plate_path(plate)
    if not path.exists():
        render_iranian_plate(plate).save(path, "PNG", optimize=True)
    return path


# ---------------------------------------------------------------------------
# Charts (blue theme)
# ---------------------------------------------------------------------------

def style_axes(ax, title: str) -> None:
    ax.set_title(title, color=BLUE["navy"], fontsize=13, fontweight="bold", pad=12)
    ax.set_facecolor(BLUE["pale"])
    ax.grid(True, axis="y", color=BLUE["sky"], linestyle="--", alpha=0.8)
    ax.tick_params(colors=BLUE["text"])
    for spine in ax.spines.values():
        spine.set_color(BLUE["mid"])


def save_fig(fig, name: str) -> Path:
    path = CHARTS / name
    fig.savefig(path, dpi=160, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def build_charts(df: pd.DataFrame) -> dict[str, Path]:
    paths: dict[str, Path] = {}

    # 1) Sessions per day
    daily = df.groupby("entry_date").size()
    fig, ax = plt.subplots(figsize=(11, 4.2))
    ax.fill_between(daily.index, daily.values, color=BLUE["light"], alpha=0.35)
    ax.plot(daily.index, daily.values, color=BLUE["primary"], linewidth=2.2, marker="o", markersize=3.5)
    style_axes(ax, "Parking Sessions per Day (Tehran Time)")
    ax.set_ylabel("Sessions")
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%m-%d"))
    fig.autofmt_xdate(rotation=35)
    paths["daily"] = save_fig(fig, "sessions_per_day.png")

    # 2) Hourly distribution
    hourly = df.groupby("entry_hour").size().reindex(range(24), fill_value=0)
    fig, ax = plt.subplots(figsize=(10, 4))
    bars = ax.bar(hourly.index, hourly.values, color=BLUE["mid"], edgecolor=BLUE["navy"], width=0.75)
    for i, b in enumerate(bars):
        if hourly.values[i] == hourly.values.max():
            b.set_color(BLUE["navy"])
    style_axes(ax, "Entry Traffic by Hour of Day (Tehran)")
    ax.set_xlabel("Hour")
    ax.set_ylabel("Entries")
    ax.set_xticks(range(0, 24, 2))
    paths["hourly"] = save_fig(fig, "entries_by_hour.png")

    # 3) Status breakdown
    status_counts = df["status"].value_counts()
    colors = [BLUE["primary"], BLUE["warn"], BLUE["success"], BLUE["danger"], BLUE["light"]]
    fig, ax = plt.subplots(figsize=(6.5, 6.5))
    wedges, texts, autotexts = ax.pie(
        status_counts.values,
        labels=status_counts.index.str.title(),
        autopct=lambda p: f"{p:.1f}%",
        colors=colors[: len(status_counts)],
        startangle=90,
        wedgeprops={"linewidth": 2, "edgecolor": "white"},
        textprops={"color": BLUE["text"]},
    )
    for t in autotexts:
        t.set_color("white")
        t.set_fontweight("bold")
    ax.set_title("Session Status Distribution", color=BLUE["navy"], fontsize=13, fontweight="bold")
    paths["status"] = save_fig(fig, "status_distribution.png")

    # 4) Payment methods
    pay = df["payment_method"].value_counts()
    fig, ax = plt.subplots(figsize=(8, 4.2))
    ax.barh(pay.index.astype(str), pay.values, color=[BLUE["primary"], BLUE["mid"], BLUE["light"], BLUE["sky"]][: len(pay)])
    style_axes(ax, "Payment Method Breakdown")
    ax.set_xlabel("Sessions")
    for i, v in enumerate(pay.values):
        ax.text(v + max(pay.values) * 0.01, i, str(v), va="center", color=BLUE["navy"], fontweight="bold")
    paths["payment"] = save_fig(fig, "payment_methods.png")

    # 5) Duration buckets
    closed = df[df["duration_minutes"].notna()].copy()
    bins = [0, 5, 15, 30, 60, 120, 240, 10**9]
    labels = ["0–5m", "5–15m", "15–30m", "30–60m", "1–2h", "2–4h", "4h+"]
    closed["bucket"] = pd.cut(closed["duration_minutes"], bins=bins, labels=labels, right=False)
    bucket_counts = closed["bucket"].value_counts().reindex(labels).fillna(0)
    fig, ax = plt.subplots(figsize=(9, 4.2))
    ax.bar(labels, bucket_counts.values, color=BLUE["accent"], edgecolor=BLUE["navy"])
    style_axes(ax, "Parking Duration Distribution")
    ax.set_ylabel("Sessions")
    paths["duration"] = save_fig(fig, "duration_distribution.png")

    # 6) Status stacked by day (last 14 active days)
    top_days = daily.sort_index().tail(14).index
    sub = df[df["entry_date"].isin(top_days)]
    pivot = pd.crosstab(sub["entry_date"], sub["status"]).reindex(top_days).fillna(0)
    fig, ax = plt.subplots(figsize=(11, 4.5))
    bottom = None
    palette = {
        "waived": BLUE["light"],
        "flagged": BLUE["warn"],
        "completed": BLUE["success"],
        "parked": BLUE["primary"],
        "unpaid": BLUE["danger"],
    }
    for col in pivot.columns:
        vals = pivot[col].values
        ax.bar(
            range(len(pivot)),
            vals,
            bottom=bottom,
            label=str(col).title(),
            color=palette.get(str(col), BLUE["mid"]),
            edgecolor="white",
            width=0.7,
        )
        bottom = vals if bottom is None else bottom + vals
    ax.set_xticks(range(len(pivot)))
    ax.set_xticklabels([d.strftime("%m-%d") for d in pivot.index], rotation=35)
    style_axes(ax, "Daily Status Mix (Recent Active Days)")
    ax.legend(frameon=False, loc="upper left")
    ax.set_ylabel("Sessions")
    paths["status_daily"] = save_fig(fig, "status_by_day.png")

    # 7) Fee analytics (non-zero)
    fee_pos = df[df["fee"] > 0]["fee"]
    fig, ax = plt.subplots(figsize=(9, 4.2))
    if len(fee_pos):
        ax.hist(fee_pos, bins=30, color=BLUE["mid"], edgecolor=BLUE["navy"], alpha=0.9)
    style_axes(ax, "Calculated Fee Distribution (Toman, fee > 0)")
    ax.set_xlabel("Fee (Toman)")
    ax.set_ylabel("Sessions")
    paths["fees"] = save_fig(fig, "fee_distribution.png")

    # 8) Entry vs exit camera coverage
    cam = pd.Series(
        {
            "Entry only": int(((df["entry_camera"] != "") & (df["exit_camera"] == "")).sum()),
            "Entry + Exit": int(((df["entry_camera"] != "") & (df["exit_camera"] != "")).sum()),
            "No camera tag": int(((df["entry_camera"] == "") & (df["exit_camera"] == "")).sum()),
        }
    )
    fig, ax = plt.subplots(figsize=(7, 4.2))
    ax.bar(cam.index, cam.values, color=[BLUE["primary"], BLUE["success"], BLUE["muted"]])
    style_axes(ax, "Camera Tag Coverage on Sessions")
    ax.set_ylabel("Sessions")
    paths["cameras"] = save_fig(fig, "camera_coverage.png")

    # Cover banner
    fig, ax = plt.subplots(figsize=(12, 2.2))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.imshow([[0, 1], [0, 1]], extent=[0, 1, 0, 1], aspect="auto", cmap="Blues", alpha=0.95)
    ax.text(0.5, 0.58, "PARKINOX", color="white", fontsize=28, fontweight="bold", ha="center", va="center")
    ax.text(0.5, 0.28, "Smart University Parking — Success & Analytics", color="#E3F2FD", fontsize=12, ha="center")
    paths["banner"] = save_fig(fig, "banner.png")

    return paths


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run_font(run, size=11, bold=False, color=None, name="Calibri") -> None:
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), "Noto Sans Arabic")
    if color:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))


def add_heading_blue(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(BLUE["navy"].lstrip("#"))


def add_body(doc: Document, text: str, size=11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, color=BLUE["text"])
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLUE["text"])


def add_kpi_table(doc: Document, kpis: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(kpis))
    table.autofit = True
    row = table.rows[0]
    for i, (label, value) in enumerate(kpis):
        cell = row.cells[i]
        set_cell_shading(cell, BLUE["pale"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(value + "\n")
        set_run_font(r1, size=16, bold=True, color=BLUE["navy"])
        r2 = p.add_run(label)
        set_run_font(r2, size=9, color=BLUE["muted"])
    doc.add_paragraph()


def add_chart(doc: Document, path: Path, width_in: float = 6.2) -> None:
    doc.add_picture(str(path), width=Inches(width_in))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)


def add_cover(doc: Document, title: str, subtitle: str, meta_lines: list[str], banner: Path) -> None:
    configure_doc_styles(doc)
    if banner.exists():
        add_chart(doc, banner, width_in=6.5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, size=26, bold=True, color=BLUE["navy"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run_font(r, size=13, color=BLUE["primary"])

    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run_font(r, size=10, color=BLUE["muted"])
    doc.add_paragraph()


def compute_stats(df: pd.DataFrame) -> dict:
    closed = df[df["exit_time"].notna()]
    return {
        "total": len(df),
        "unique_plates": int(df["plate"].nunique()),
        "completed": int((df["status"] == "completed").sum()),
        "waived": int((df["status"] == "waived").sum()),
        "flagged": int((df["status"] == "flagged").sum()),
        "wallet": int((df["payment_method"] == "wallet").sum()),
        "cash": int((df["payment_method"] == "cash").sum()),
        "with_exit": int(df["exit_time"].notna().sum()),
        "open_or_flagged_no_exit": int(df["exit_time"].isna().sum()),
        "fee_sum": int(df["fee"].sum()),
        "fee_mean_pos": float(df.loc[df["fee"] > 0, "fee"].mean()) if (df["fee"] > 0).any() else 0.0,
        "avg_duration": float(closed["duration_minutes"].mean()) if len(closed) else 0.0,
        "median_duration": float(closed["duration_minutes"].median()) if len(closed) else 0.0,
        "date_min": df["entry_tehran"].min(),
        "date_max": df["entry_tehran"].max(),
        "peak_day": df.groupby("entry_date").size().idxmax(),
        "peak_day_count": int(df.groupby("entry_date").size().max()),
        "peak_hour": int(df.groupby("entry_hour").size().idxmax()),
        "registered_named": int((df["user_name"] != "").sum()),
    }


# ---------------------------------------------------------------------------
# Document 1 — Success Report
# ---------------------------------------------------------------------------

def build_success_report(df: pd.DataFrame, charts: dict[str, Path], stats: dict) -> Path:
    doc = Document()
    add_cover(
        doc,
        "Parkinox Success Report",
        "End-to-End Smart Parking Platform — Operational Success & Session Ledger",
        [
            "Blue Theme Edition · Iranian Plate Recognition · Local-First Architecture",
            f"Data window: {fmt_dt(stats['date_min'])}  →  {fmt_dt(stats['date_max'])} (Tehran)",
            f"Generated: {datetime.now(TEHRAN).strftime('%Y-%m-%d %H:%M')} Asia/Tehran",
            "Source: parkinox_sessions.csv · Project: Parkinox v3",
        ],
        charts["banner"],
    )

    add_heading_blue(doc, "1. Executive Summary", 1)
    add_body(
        doc,
        "Parkinox v3 successfully delivered a privacy-preserving university parking platform that "
        "combines local Iranian plate recognition (YOLOv5 on the operator PC), Django business logic "
        "for sessions/fees/wallets, Flutter operator and student clients, and a durable outbox sync "
        "to a remote reporting cloud. Live operations in the attached session export demonstrate "
        f"{stats['total']:,} recorded parking sessions across {stats['unique_plates']:,} unique plates, "
        "with gate confirmation, fee calculation, payment paths, and operator waive/flag workflows "
        "functioning as designed.",
    )

    add_kpi_table(
        doc,
        [
            ("Total Sessions", f"{stats['total']:,}"),
            ("Unique Plates", f"{stats['unique_plates']:,}"),
            ("Completed", f"{stats['completed']:,}"),
            ("Waived", f"{stats['waived']:,}"),
            ("Flagged / Open", f"{stats['flagged']:,}"),
        ],
    )

    add_heading_blue(doc, "2. How the System Functions", 1)
    add_heading_blue(doc, "2.1 Architecture", 2)
    add_body(
        doc,
        "Detection stays on the operator machine: FastAPI + OpenCV + YOLOv5 ingest RTSP/USB cameras, "
        "emit plate_detected events, and serve MJPEG preview. The remote Django server is the source of "
        "truth for authentication, parking sessions, gate events, wallets, and reports. Student devices "
        "use the Flutter mobile app (OTP, plate registration, wallet). A separate cloud Django replica "
        "receives metadata/images via SyncOutbox for analytics — never raw camera video.",
    )
    add_bullet(doc, "parkinox_op — Flutter desktop operator UI (approval, gates, fail review)")
    add_bullet(doc, "Core/FastAPI — plateYolo.pt + CharsYolo.pt, validation, success/fail capture")
    add_bullet(doc, "backend/Django — GateEvent, ParkingSession, RateConfig, wallets, reports")
    add_bullet(doc, "parkinox_stdn — student mobile access")
    add_bullet(doc, "cloud/ — Persian RTL reporting dashboard (Chart.js, Excel/PDF/CSV)")

    add_heading_blue(doc, "2.2 Session Lifecycle", 2)
    add_body(
        doc,
        "Camera frame → YOLO plate bbox → character OCR → Iranian format validation (same normalizer "
        "as Django) → operator auto-approve countdown or manual confirm → confirm_gate_event():",
    )
    add_bullet(doc, "Entry: create ParkingSession (status=parked); link immutable GateEvent")
    add_bullet(doc, "Exit: close session, compute duration/fee (grace, hourly rate, daily cap, student discount)")
    add_bullet(doc, "Payment: wallet auto-deduct when registered & funded; else unpaid → cash collect or waive")
    add_bullet(doc, "Post-commit: SyncOutbox enqueue + success-capture finalize (scene/crop media)")

    add_heading_blue(doc, "2.3 Iranian Plate Handling", 2)
    add_body(
        doc,
        "National plates use canonical form ۱۲ب۳۴۵-۶۷; free-zone plates use ۵+۲ digits. Normalization "
        "converts Arabic-Indic/Latin digits, strips noise, and validates the legal letter set. Operator "
        "display spaces the plate (۱۲ ب ۳۴۵ ۶۷). Cross-camera guard blocks the opposite direction for "
        "the same plate for two minutes to prevent bounce errors.",
    )

    # Sample plates showcase
    sample = df.head(6)
    add_body(doc, "Sample plates rendered in Iranian plate style (from operational export):")
    table = doc.add_table(rows=1, cols=3)
    for i, (_, row) in enumerate(sample.iterrows()):
        if i >= 3:
            break
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, BLUE["pale"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img = get_plate_image(row["plate"])
        run = p.add_run()
        run.add_picture(str(img), width=Inches(1.85))
    doc.add_paragraph()

    add_heading_blue(doc, "3. Error Handling & Resilience", 1)
    add_heading_blue(doc, "3.1 Detection Fail Capture", 2)
    add_body(
        doc,
        "Failed OCR never auto-opens a gate. Failures are classified and stored for operator review:",
    )
    add_bullet(doc, "Type A — ocr_invalid_format: bbox found, text fails Iranian rules")
    add_bullet(doc, "Type B — ocr_empty: bbox found, empty/low-confidence OCR")
    add_bullet(doc, "Type C — operator_confirmed_miss: visible plate, no bbox (operator capture)")
    add_bullet(doc, "Type D — transient retry success metric (fail then valid within ~5s)")
    add_body(
        doc,
        "Core FailCaptureStore applies cooldown, per-camera rate limits, IoU dedup, and best-in-window "
        "selection before ingest to Django DetectionFail. Operators can correct the plate (calls "
        "confirm_gate_event with original captured_at) or mark not_a_plate.",
    )

    add_heading_blue(doc, "3.2 Gate & Session Integrity", 2)
    add_bullet(doc, "GateEvent is an immutable audit log (insert-only)")
    add_bullet(doc, "Flags: was_edited, was_auto_approved, was_rejected")
    add_bullet(doc, "Concurrent parked sessions for the same plate are flagged/auto-closed before re-entry")
    add_bullet(doc, "Historical fail corrections are rejected if they would corrupt a newer live session")
    add_bullet(doc, "Wallet deduct uses select_for_update; insufficient funds → unpaid (no partial charge)")

    add_heading_blue(doc, "3.3 Pipeline & Sync Resilience", 2)
    add_bullet(doc, "Camera workers: reconnect backoff, blur filter, detection interval, emit cooldowns")
    add_bullet(doc, "Success capture buffers frames; finalize is non-blocking (buffer_miss logged if TTL expired)")
    add_bullet(doc, "SyncOutbox enqueued only after DB commit — never on the gate critical path")
    add_bullet(doc, "Outbox statuses: pending → sending → acked / dead_letter (max 12 attempts)")
    add_bullet(doc, "Local gates and reporting keep working during cloud outage")

    add_heading_blue(doc, "4. Operational Outcomes from Session Data", 1)
    add_body(
        doc,
        f"Across the export window, peak day {stats['peak_day']} recorded {stats['peak_day_count']:,} "
        f"sessions. Peak entry hour (Tehran) was {stats['peak_hour']:02d}:00. "
        f"{stats['with_exit']:,} sessions have an exit timestamp; {stats['open_or_flagged_no_exit']:,} "
        f"remain without exit (typically flagged). Payment methods observed: wallet={stats['wallet']}, "
        f"cash={stats['cash']}. Operator waive workflow accounts for {stats['waived']:,} closures — "
        "consistent with pilot/campus policy use rather than a system fault.",
    )
    add_kpi_table(
        doc,
        [
            ("Avg Duration", fmt_duration(stats["avg_duration"])),
            ("Median Duration", fmt_duration(stats["median_duration"])),
            ("Fee Sum (Tomans)", f"{stats['fee_sum']:,}"),
            ("Named Users", f"{stats['registered_named']:,}"),
            ("Peak Hour", f"{stats['peak_hour']:02d}:00"),
        ],
    )

    add_heading_blue(doc, "5. Analytics Charts", 1)
    add_body(doc, "Charts below are generated from the same CSV used in the companion Data Analytics Report.")
    for key, caption in [
        ("daily", "Figure 1 — Sessions per day"),
        ("hourly", "Figure 2 — Entry traffic by hour"),
        ("status", "Figure 3 — Status distribution"),
        ("payment", "Figure 4 — Payment methods"),
        ("duration", "Figure 5 — Duration buckets"),
        ("status_daily", "Figure 6 — Recent daily status mix"),
        ("fees", "Figure 7 — Fee distribution"),
        ("cameras", "Figure 8 — Camera tag coverage"),
    ]:
        add_chart(doc, charts[key])
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        set_run_font(r, size=9, color=BLUE["muted"])

    add_heading_blue(doc, "6. Complete Session Ledger — All Plates", 1)
    add_body(
        doc,
        "All sessions from the CSV, sorted by entry time (newest first). Each row shows the Iranian "
        "plate badge, status, Tehran entry/exit times, duration, fee, and payment method.",
    )

    # Compact ledger table — plate image + meta
    # Use batches of header + rows for readability
    header_labels = ["#", "Plate", "Status", "Entry (Tehran)", "Exit (Tehran)", "Duration", "Fee", "Pay"]
    batch_size = 40
    total = len(df)
    for start in range(0, total, batch_size):
        chunk = df.iloc[start : start + batch_size]
        if start > 0:
            add_heading_blue(doc, f"Sessions {start + 1}–{min(start + batch_size, total)}", 2)
        else:
            add_heading_blue(doc, f"Sessions 1–{min(batch_size, total)}", 2)

        table = doc.add_table(rows=1, cols=len(header_labels))
        table.style = "Table Grid"
        hdr = table.rows[0]
        for i, label in enumerate(header_labels):
            cell = hdr.cells[i]
            set_cell_shading(cell, BLUE["navy"])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(label)
            set_run_font(r, size=8, bold=True, color=BLUE["white"])

        for idx, (_, row) in enumerate(chunk.iterrows(), start=start + 1):
            cells = table.add_row().cells
            values = [
                str(idx),
                "",  # plate image
                str(row["status"]),
                fmt_dt(row["entry_tehran"]),
                fmt_dt(row["exit_tehran"]),
                fmt_duration(row["duration_minutes"]),
                f"{int(row['fee']):,}" if row["fee"] else "0",
                str(row["payment_method"]),
            ]
            for i, val in enumerate(values):
                p = cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i == 1:
                    img = get_plate_image(row["plate"])
                    run = p.add_run()
                    run.add_picture(str(img), width=Inches(1.35))
                else:
                    r = p.add_run(val)
                    set_run_font(r, size=7, color=BLUE["text"])
                if idx % 2 == 0:
                    set_cell_shading(cells[i], BLUE["pale"])

        doc.add_paragraph()

    add_heading_blue(doc, "7. Conclusion", 1)
    add_body(
        doc,
        "Parkinox met its success criteria for the pilot: local Iranian plate recognition with "
        "validated OCR, operator-controlled gate confirmation, durable session accounting, wallet/cash/"
        "waive payment paths, structured fail capture without unsafe auto-open, and cloud reporting "
        "that never compromises the local source of truth. The session ledger and charts in this "
        "document provide an auditable record of system behavior under real campus traffic.",
    )

    out = ROOT / "Parkinox_Success_Report.docx"
    doc.save(out)
    return out


# ---------------------------------------------------------------------------
# Document 2 — Data Analytics only
# ---------------------------------------------------------------------------

def build_data_report(df: pd.DataFrame, charts: dict[str, Path], stats: dict) -> Path:
    doc = Document()
    add_cover(
        doc,
        "Parkinox Data Analytics Report",
        "Session Dataset Charts & Measured Date/Time Insights",
        [
            "Blue Theme · Data-Only Document (CSV Driven)",
            f"Rows: {stats['total']:,}  ·  Unique plates: {stats['unique_plates']:,}",
            f"Window: {fmt_dt(stats['date_min'])}  →  {fmt_dt(stats['date_max'])} (Tehran)",
            f"Generated: {datetime.now(TEHRAN).strftime('%Y-%m-%d %H:%M')} Asia/Tehran",
            "Source file: parkinox_sessions.csv",
        ],
        charts["banner"],
    )

    add_heading_blue(doc, "1. Dataset Overview", 1)
    add_body(
        doc,
        "This report contains only measurements derived from the provided parking sessions CSV. "
        "Timestamps were converted from UTC to Asia/Tehran (UTC+03:30) for traffic analysis. "
        "Duration and fee fields are taken as exported; sessions without exit_time are treated as open/flagged.",
    )
    add_kpi_table(
        doc,
        [
            ("Sessions", f"{stats['total']:,}"),
            ("Unique Plates", f"{stats['unique_plates']:,}"),
            ("With Exit", f"{stats['with_exit']:,}"),
            ("No Exit", f"{stats['open_or_flagged_no_exit']:,}"),
            ("Fee Total", f"{stats['fee_sum']:,}"),
        ],
    )

    # Column schema
    add_heading_blue(doc, "2. Schema", 1)
    schema = [
        ("plate", "Iranian plate string (Persian digits + letter)"),
        ("status", "completed | waived | flagged"),
        ("entry_time / exit_time", "ISO-8601 UTC timestamps"),
        ("duration_minutes", "Measured stay length when exited"),
        ("fee", "Calculated fee in Tomans"),
        ("payment_method", "wallet | cash | waived | none"),
        ("entry_camera / exit_camera", "Camera tags when present"),
        ("user_name", "Linked user display name when registered"),
        ("session_uuid", "Stable session identifier"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, h in enumerate(["Field", "Meaning"]):
        set_cell_shading(table.rows[0].cells[i], BLUE["navy"])
        r = table.rows[0].cells[i].paragraphs[0].add_run(h)
        set_run_font(r, size=10, bold=True, color=BLUE["white"])
    for field, meaning in schema:
        row = table.add_row().cells
        r0 = row[0].paragraphs[0].add_run(field)
        set_run_font(r0, size=9, bold=True, color=BLUE["navy"])
        r1 = row[1].paragraphs[0].add_run(meaning)
        set_run_font(r1, size=9, color=BLUE["text"])
    doc.add_paragraph()

    add_heading_blue(doc, "3. Status & Payment Counts", 1)
    status = df["status"].value_counts()
    pay = df["payment_method"].value_counts()
    add_body(doc, "Status counts:")
    for k, v in status.items():
        add_bullet(doc, f"{k}: {v:,} ({100 * v / len(df):.1f}%)")
    add_body(doc, "Payment method counts:")
    for k, v in pay.items():
        add_bullet(doc, f"{k}: {v:,} ({100 * v / len(df):.1f}%)")

    add_heading_blue(doc, "4. Time-Series Charts", 1)
    add_body(
        doc,
        f"Peak day: {stats['peak_day']} with {stats['peak_day_count']:,} entries. "
        f"Peak hour (Tehran): {stats['peak_hour']:02d}:00. "
        f"Average closed duration: {fmt_duration(stats['avg_duration'])}; "
        f"median: {fmt_duration(stats['median_duration'])}.",
    )
    for key, caption in [
        ("daily", "Chart A — Sessions per day"),
        ("hourly", "Chart B — Entries by hour of day"),
        ("status_daily", "Chart C — Status mix on recent active days"),
    ]:
        add_chart(doc, charts[key])
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        set_run_font(r, size=9, color=BLUE["muted"])

    add_heading_blue(doc, "5. Composition Charts", 1)
    for key, caption in [
        ("status", "Chart D — Status distribution"),
        ("payment", "Chart E — Payment methods"),
        ("duration", "Chart F — Duration distribution"),
        ("fees", "Chart G — Fee distribution (fee > 0)"),
        ("cameras", "Chart H — Camera tag coverage"),
    ]:
        add_chart(doc, charts[key])
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        set_run_font(r, size=9, color=BLUE["muted"])

    add_heading_blue(doc, "6. Daily Traffic Table", 1)
    daily = df.groupby("entry_date").agg(
        sessions=("plate", "size"),
        unique_plates=("plate", "nunique"),
        completed=("status", lambda s: (s == "completed").sum()),
        waived=("status", lambda s: (s == "waived").sum()),
        flagged=("status", lambda s: (s == "flagged").sum()),
        fee_sum=("fee", "sum"),
    )
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["Date", "Sessions", "Unique", "Completed", "Waived", "Flagged", "Fee Sum"]
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], BLUE["primary"])
        r = table.rows[0].cells[i].paragraphs[0].add_run(h)
        set_run_font(r, size=8, bold=True, color=BLUE["white"])
    for date, row in daily.iterrows():
        cells = table.add_row().cells
        vals = [
            str(date),
            f"{int(row['sessions']):,}",
            f"{int(row['unique_plates']):,}",
            f"{int(row['completed']):,}",
            f"{int(row['waived']):,}",
            f"{int(row['flagged']):,}",
            f"{int(row['fee_sum']):,}",
        ]
        for i, v in enumerate(vals):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(v)
            set_run_font(r, size=8, color=BLUE["text"])
    doc.add_paragraph()

    add_heading_blue(doc, "7. Hourly Traffic Table", 1)
    hourly = df.groupby("entry_hour").size().reindex(range(24), fill_value=0)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, h in enumerate(["Hour (Tehran)", "Entries"]):
        set_cell_shading(table.rows[0].cells[i], BLUE["primary"])
        r = table.rows[0].cells[i].paragraphs[0].add_run(h)
        set_run_font(r, size=9, bold=True, color=BLUE["white"])
    for hour, count in hourly.items():
        cells = table.add_row().cells
        for i, v in enumerate([f"{int(hour):02d}:00", f"{int(count):,}"]):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(v)
            set_run_font(r, size=9, color=BLUE["text"])
            if count == hourly.max() and count > 0:
                set_cell_shading(cells[i], BLUE["sky"])
    doc.add_paragraph()

    add_heading_blue(doc, "8. Duration & Fee Statistics", 1)
    closed = df[df["duration_minutes"].notna()]
    add_bullet(doc, f"Closed sessions: {len(closed):,}")
    add_bullet(doc, f"Duration mean: {fmt_duration(closed['duration_minutes'].mean())}")
    add_bullet(doc, f"Duration median: {fmt_duration(closed['duration_minutes'].median())}")
    add_bullet(doc, f"Duration p90: {fmt_duration(closed['duration_minutes'].quantile(0.9))}")
    add_bullet(doc, f"Duration max: {fmt_duration(closed['duration_minutes'].max())}")
    add_bullet(doc, f"Total calculated fees: {stats['fee_sum']:,} Tomans")
    add_bullet(doc, f"Mean fee where fee > 0: {stats['fee_mean_pos']:,.0f} Tomans")
    add_bullet(doc, f"Sessions with fee = 0: {int((df['fee'] == 0).sum()):,}")

    add_heading_blue(doc, "9. Notes on Measurement", 1)
    add_body(
        doc,
        "Entry and exit instants are parsed from ISO timestamps with timezone offsets, then converted "
        "to Tehran local time for hourly/daily aggregation. Duration uses the CSV duration_minutes "
        "field (not re-derived) to stay faithful to the export. Status and payment_method are "
        "categorical counts; empty payment_method values are labeled none. This document intentionally "
        "excludes product narrative — see Parkinox_Success_Report.docx for system function and error handling.",
    )

    out = ROOT / "Parkinox_Data_Analytics_Report.docx"
    doc.save(out)
    return out


def main() -> None:
    ensure_dirs()
    print("Loading CSV…")
    df = load_data()
    print(f"Loaded {len(df)} sessions, {df['plate'].nunique()} unique plates")

    print("Pre-rendering unique plate badges…")
    for plate in df["plate"].unique():
        get_plate_image(plate)
    print(f"Plate images in {PLATES}")

    print("Building English charts…")
    charts = build_charts(df)
    stats = compute_stats(df)

    print("Writing English Success Report DOCX…")
    success = build_success_report(df, charts, stats)
    print(f"  → {success}")

    print("Writing English Data Analytics DOCX…")
    data = build_data_report(df, charts, stats)
    print(f"  → {data}")

    # Persian (RTL) editions
    from persian_reports import (
        build_charts_fa,
        build_data_report_fa,
        build_success_report_fa,
    )

    print("Building Persian charts…")
    charts_fa = build_charts_fa(df)

    print("Writing Persian Success Report DOCX…")
    success_fa = build_success_report_fa(df, charts_fa, stats)
    print(f"  → {success_fa}")

    print("Writing Persian Data Analytics DOCX…")
    data_fa = build_data_report_fa(df, charts_fa, stats)
    print(f"  → {data_fa}")

    # Write a short README for the report folder
    readme = ROOT / "README.md"
    readme.write_text(
        "# Parkinox Reports\n\n"
        "Blue-themed success and data analytics documents (English + Persian) "
        "generated from the operational session export.\n\n"
        "## Outputs\n\n"
        "| File | Description |\n"
        "|------|-------------|\n"
        "| `Parkinox_Success_Report.docx` | English success narrative, architecture, error handling, Iranian plate ledger, charts |\n"
        "| `Parkinox_Success_Report_FA.docx` | **Persian RTL** edition of the success report |\n"
        "| `Parkinox_Data_Analytics_Report.docx` | English CSV-only analytics |\n"
        "| `Parkinox_Data_Analytics_Report_FA.docx` | **Persian RTL** edition of the data analytics report |\n"
        "| `charts/` | PNG charts (English + `fa_*` Persian-labeled) |\n"
        "| `plates/` | Cached Iranian-style plate badge PNGs |\n"
        "| `parkinox_sessions.csv` | Source session export |\n"
        "| `generate_reports.py` / `persian_reports.py` | Generators |\n\n"
        "## Regenerate\n\n"
        "```bash\n"
        "pip install python-docx matplotlib pillow pandas\n"
        "python report/generate_reports.py\n"
        "```\n",
        encoding="utf-8",
    )
    print("Done.")


if __name__ == "__main__":
    main()
